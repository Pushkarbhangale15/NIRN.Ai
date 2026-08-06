"""
routes.py — every endpoint in NIRN.Ai.

SQL RULE: this file never builds a query itself. Every draft/officer
read or write goes through db.repositories.*, which is the only place
that talks SQLAlchemy. Never an f-string, never string concatenation.
See backend/README.md, "SQL injection prevention".

Coming from Express, the mapping is:
    APIRouter()                 ->  express.Router()
    @router.post("/x")          ->  router.post("/x", handler)
    HTTPException(404, detail)  ->  res.status(404).json({...})

The difference in your favour: FastAPI validates the request body
against the schema automatically and returns 422 with a field-by-field
explanation, so you never write manual field checks.
"""

import difflib
import io
import logging
import re
import time
import hashlib
import uuid
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, status, File, UploadFile
from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.ext.asyncio import AsyncSession

import llm
import prompts
import references
import retrieval
import store
import template_rules
import template_rules_marathi
from conflict_detection import detect_cross_department_conflicts
from ocr_ingest.pipeline import run_ocr_pipeline
from ocr_ingest.storage import save_upload_bytes
from conflict_detection.llm_verifier import verify_conflict_with_llm
from conflict_detection.models import ClauseRetrievalTrace, ConflictReportItem
from lookup import get_adapter
from config import settings
from db.base import get_session
from db.base import engine as db_engine
from db.models import Officer, OfficerRole
from db.repositories import conflicts as conflicts_repo
from db.repositories import drafts as drafts_repo
from db.repositories import gr_uploads as gr_uploads_repo
from deps import get_current_officer, optional_auth
from schemas import (
    Language,
    AnalysisReport,
    AnalysisSummary,
    AcceptConflictResolutionRequest,
    AcceptConflictResolutionResponse,
    ConflictHit,
    ConflictOut,
    DismissConflictRequest,
    GrUploadResponse,
    ResolveConflictRequest,
    ResolveConflictResponse,
    ReverificationResult,
    DraftReferenceOut,
    GeneratedDraftDetail,
    DraftHistoryItem,
    HealthDbResponse,
    PaginatedDraftHistory,
    Relation,
    CorpusSearchResponse,
    FullOCRResponse,
    OfficialSourceResponse,
    Draft,
    DraftCreate,
    DraftUpdate,
    ReferenceHit,
    Severity,
    TemplateIssue,
    TermMapping,
    DraftGenerateRequest,
    DraftGenerateResponse,
    ComparisonRequest,
    ComparisonResponse,
    ClauseExplanationRequest,
    ClauseExplanationResponse,
)

logger = logging.getLogger("nirn.routes")

router = APIRouter()

_DEVANAGARI_RE = re.compile(r"[ऀ-ॿ]")

# The LLM verifier's severity scale (Low/Medium/High/Critical) is wider than
# the DB enum (low/medium/high) — Critical has no DB counterpart, so it maps
# down to "high" rather than failing the insert.
_DB_SEVERITY_MAP = {"low": "low", "medium": "medium", "high": "high", "critical": "high"}


def _to_db_severity(severity: Optional[str]) -> str:
    return _DB_SEVERITY_MAP.get((severity or "").lower(), "medium")


# =====================================================================
# Draft (de)serialisation helpers — map between the GeneratedDraft ORM
# row and the pydantic shapes the frontend already speaks.
# =====================================================================

def _to_draft_schema(row) -> Draft:
    return Draft(
        id=str(row.generated_draft_id),
        title=row.title,
        department=row.department,
        body_text=row.content,
        language=row.language.value if hasattr(row.language, "value") else row.language,
        created_at=row.created_at,
        gr_number=row.gr_number,
        status=row.status.value if hasattr(row.status, "value") else row.status,
        returned_reason=row.returned_reason,
    )


def _to_detail_schema(row) -> GeneratedDraftDetail:
    return GeneratedDraftDetail(
        generated_draft_id=row.generated_draft_id,
        gr_number=row.gr_number,
        title=row.title,
        department=row.department,
        language=row.language.value if hasattr(row.language, "value") else row.language,
        status=row.status.value if hasattr(row.status, "value") else row.status,
        version=row.version,
        content=row.content,
        content_plain=row.content_plain,
        brief=row.brief,
        drafted_by=row.drafted_by,
        drafted_by_name=row.officer.name if row.officer is not None else None,
        created_at=row.created_at,
        updated_at=row.updated_at,
        returned_reason=row.returned_reason,
        conflicts=[ConflictOut.model_validate(c) for c in row.conflicts],
        references=[DraftReferenceOut.model_validate(r) for r in row.references],
    )


def _ensure_can_access_draft(draft_row, officer: Officer) -> None:
    if officer.role in (OfficerRole.ADMIN, OfficerRole.REVIEWER):
        return
    if draft_row.drafted_by != officer.officer_id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="You do not have access to this draft.")


async def _load_row(draft_id: uuid.UUID, session: AsyncSession):
    """Fetch a draft ORM row or raise a clean 404."""
    try:
        row = await drafts_repo.get_draft_by_id(session, draft_id)
    except SQLAlchemyError:
        logger.exception("Database error loading draft %s", draft_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    if row is None:
        raise HTTPException(status_code=404, detail=f"Draft {draft_id} not found")
    return row


async def _load(draft_id: uuid.UUID, session: AsyncSession, officer: Optional[Officer] = None) -> Draft:
    """Fetch a draft or raise a clean 404. Used by every analysis route."""
    row = await _load_row(draft_id, session)
    if officer is not None:
        _ensure_can_access_draft(row, officer)
    return _to_draft_schema(row)


def _severity_to_db(sev: Optional[str]) -> str:
    s = (sev or "medium").strip().lower()
    return s if s in ("low", "medium", "high") else "medium"


def _conflict_hit_to_dict(c: ConflictHit) -> dict:
    return {
        "source_of_conflict": c.existing_department or c.existing_gr_title or "Unknown",
        "conflicting_text": c.existing_clause,
        "draft_excerpt": c.draft_clause,
        "conflicting_gr_id": c.existing_gr_id,
        "source_gr_title": c.existing_gr_title,
        "severity": _severity_to_db(c.severity),
        "justification": c.justification,
        "detected_by": "llm_verifier",
    }


def _reference_hit_to_dict(r: ReferenceHit) -> dict:
    return {
        "reference_text": r.raw_text,
        "extracted_gr_number": r.gr_number,
        "script": "devanagari" if _DEVANAGARI_RE.search(r.raw_text) else "latin",
        "resolved": r.found_in_corpus,
    }


# =====================================================================
# Drafts — CRUD + history (Task 6)
# =====================================================================

@router.post("/api/drafts", response_model=Draft,
             status_code=status.HTTP_201_CREATED, tags=["drafts"])
async def create_draft(
    payload: DraftCreate,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> Draft:
    """Submit a draft GR. Requires login — this is a "generating a draft" action."""
    try:
        draft = await drafts_repo.create_draft_with_analysis(
            session,
            title=payload.title,
            language=payload.language.value,
            drafted_by=officer.officer_id,
            content=payload.body_text,
            content_plain=payload.body_text,
            department=payload.department,
            brief=None,
        )
    except SQLAlchemyError:
        logger.exception("Database error creating draft")
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    return _to_draft_schema(draft)


@router.get("/api/drafts", response_model=PaginatedDraftHistory, tags=["drafts"])
async def list_drafts(
    status_filter: Optional[str] = Query(None, alias="status"),
    department: Optional[str] = Query(None, max_length=160),
    search: Optional[str] = Query(None, max_length=200),
    sort_by: str = Query("created_at"),
    sort_desc: bool = Query(True),
    page: int = Query(1, ge=1),
    page_size: int = Query(default=None, ge=1, le=100),
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> PaginatedDraftHistory:
    """History listing. An officer sees only their own drafts; a reviewer/admin sees all."""
    size = min(page_size or settings.DEFAULT_PAGE_SIZE, settings.MAX_PAGE_SIZE)
    try:
        rows, total = await drafts_repo.list_drafts(
            session,
            requesting_officer_id=officer.officer_id,
            requesting_role=officer.role,
            page=page,
            page_size=size,
            status=status_filter,
            department=department,
            search=search,
            sort_by=sort_by,
            sort_desc=sort_desc,
        )
    except SQLAlchemyError:
        logger.exception("Database error listing drafts")
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")

    items = [
        DraftHistoryItem(
            generated_draft_id=d.generated_draft_id,
            gr_number=d.gr_number,
            title=d.title,
            department=d.department,
            language=d.language.value if hasattr(d.language, "value") else d.language,
            status=d.status.value if hasattr(d.status, "value") else d.status,
            version=d.version,
            created_at=d.created_at,
            updated_at=d.updated_at,
            unresolved_conflict_count=count,
            returned_reason=d.returned_reason,
        )
        for d, count in rows
    ]
    return PaginatedDraftHistory(items=items, total=total, page=page, page_size=size)


@router.get("/api/drafts/{draft_id}", response_model=GeneratedDraftDetail, tags=["drafts"])
async def get_draft(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> GeneratedDraftDetail:
    row = await _load_row(draft_id, session)
    _ensure_can_access_draft(row, officer)
    return _to_detail_schema(row)


@router.get("/api/drafts/{draft_id}/conflicts", response_model=List[ConflictOut], tags=["drafts"])
async def get_draft_conflicts(
    draft_id: uuid.UUID,
    include_dismissed: bool = Query(True),
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> List[ConflictOut]:
    row = await _load_row(draft_id, session)
    _ensure_can_access_draft(row, officer)
    try:
        rows = await conflicts_repo.get_conflicts_for_draft(session, draft_id, include_dismissed=include_dismissed)
    except SQLAlchemyError:
        logger.exception("Database error loading conflicts for draft %s", draft_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    return [ConflictOut.model_validate(c) for c in rows]


@router.get(
    "/api/drafts/{draft_id}/retrieval-trace",
    response_model=List[ClauseRetrievalTrace],
    tags=["drafts"],
)
async def get_draft_retrieval_trace(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> List[ClauseRetrievalTrace]:
    """
    Observability endpoint: per-clause retrieval detail (top_k used, every
    candidate's GR id/department/score, whether it reached the LLM or was
    filtered out earlier) from the most recent conflict-detection run
    against this draft. Present even when 0 conflicts were found — that's
    the whole point: distinguishes "checked N candidates, none conflicted"
    from "no relevant candidates were ever retrieved" for a clause.

    Empty list if analysis (POST /api/analysis/{draft_id} or
    /api/analysis/{draft_id}/conflicts) hasn't been run yet for this draft.
    """
    row = await _load_row(draft_id, session)
    _ensure_can_access_draft(row, officer)
    return [ClauseRetrievalTrace.model_validate(c) for c in (row.retrieval_trace or [])]


@router.patch("/api/conflicts/{conflict_id}/dismiss", response_model=ConflictOut, tags=["drafts"])
async def dismiss_conflict(
    conflict_id: uuid.UUID,
    payload: DismissConflictRequest,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> ConflictOut:
    try:
        conflict = await conflicts_repo.get_by_id(session, conflict_id)
    except SQLAlchemyError:
        logger.exception("Database error loading conflict %s", conflict_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    if conflict is None:
        raise HTTPException(status_code=404, detail="Conflict not found")

    draft_row = await _load_row(conflict.generated_draft_id, session)
    _ensure_can_access_draft(draft_row, officer)

    updated = await conflicts_repo.dismiss_conflict(session, conflict_id, reason=payload.reason)
    return ConflictOut.model_validate(updated)


@router.post("/api/conflicts/{conflict_id}/resolve", response_model=ResolveConflictResponse, tags=["drafts"])
async def resolve_conflict(
    conflict_id: uuid.UUID,
    payload: ResolveConflictRequest,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> ResolveConflictResponse:
    """
    Narrow, single-clause follow-up action: revise just the flagged clause
    per the chosen strategy, then re-run only that clause through the
    existing conflict-check logic (one LLM call, not the full batch) to
    confirm the revision actually clears the conflict. A successful accept
    is the only thing that patches the draft — but the outcome of every
    /resolve attempt (cleared, still-conflicting, or errored) is persisted
    as resolution_status here, so a retry after a page reload can tell
    whether this conflict was already dealt with instead of blindly
    regenerating against stale state.
    """
    logger.info("Resolve requested for conflict %s (strategy=%s)", conflict_id, payload.strategy.value)
    try:
        conflict = await conflicts_repo.get_by_id(session, conflict_id)
    except SQLAlchemyError:
        logger.exception("Database error loading conflict %s", conflict_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    if conflict is None:
        raise HTTPException(status_code=404, detail="Conflict not found")

    draft_row = await _load_row(conflict.generated_draft_id, session)
    _ensure_can_access_draft(draft_row, officer)

    # Already resolved (durable, DB-backed) — short-circuit rather than
    # regenerating: the original clause text was already replaced in the
    # draft by the prior acceptance, so re-deriving a revision from
    # conflict.draft_excerpt (stale) and re-verifying would be wasted work
    # at best and a guaranteed accept-time 409 at worst.
    if conflict.resolution_status == "resolved" and conflict.resolved_clause_text:
        logger.info("Conflict %s already resolved; returning persisted resolution", conflict_id)
        return ResolveConflictResponse(
            conflict_id=conflict_id,
            strategy=payload.strategy,
            original_clause=conflict.draft_excerpt or "",
            revised_clause=conflict.resolved_clause_text,
            diff="",
            reverification=ReverificationResult(conflict=False, relation=Relation.UNRELATED),
            cleared=True,
        )

    original_clause = conflict.draft_excerpt or ""
    if not original_clause.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="This conflict has no stored draft excerpt to revise.",
        )

    conflicting_gr_label = f"{conflict.conflicting_gr_id or 'Unknown'}: {conflict.source_gr_title or ''}".strip()
    user_msg = prompts.build_conflict_resolution_message(
        strategy=payload.strategy.value,
        draft_clause=original_clause,
        conflicting_clause=conflict.conflicting_text,
        conflicting_gr_label=conflicting_gr_label,
        justification=conflict.justification,
    )
    try:
        revised_clause = llm.call_model(prompts.CONFLICT_RESOLUTION, user_msg).strip()
        logger.info("Conflict %s: revision generated (%d chars)", conflict_id, len(revised_clause))
    except Exception:
        logger.exception("Conflict %s: revision generation failed", conflict_id)
        try:
            await conflicts_repo.record_resolve_attempt(session, conflict_id, status="attempted_error")
        except SQLAlchemyError:
            logger.exception("Database error recording failed resolve attempt for %s", conflict_id)
        raise

    diff = "\n".join(
        difflib.unified_diff(
            original_clause.splitlines(),
            revised_clause.splitlines(),
            lineterm="",
            fromfile="original", tofile="revised",
        )
    )

    # Single Gemma call to confirm the revision actually clears the conflict.
    reverify_item = verify_conflict_with_llm(
        draft_clause=revised_clause,
        matched_gr_id=conflict.conflicting_gr_id or "",
        matched_gr_title=conflict.source_gr_title or "",
        matched_clause=conflict.conflicting_text,
        matched_department=conflict.source_of_conflict,
    )
    if reverify_item is None:
        reverification = ReverificationResult(conflict=False, relation=Relation.UNRELATED)
        cleared = True
    else:
        relation = Relation.OVERLAP if reverify_item.relation == "overlap" else Relation.CONFLICT
        reverification = ReverificationResult(
            conflict=True,
            relation=relation,
            severity=reverify_item.severity,
            confidence=reverify_item.confidence,
            justification=reverify_item.reason,
        )
        cleared = relation != Relation.CONFLICT

    logger.info("Conflict %s: reverification cleared=%s", conflict_id, cleared)

    if not cleared:
        try:
            await conflicts_repo.record_resolve_attempt(
                session, conflict_id, status="attempted_still_conflicting"
            )
        except SQLAlchemyError:
            logger.exception("Database error recording still-conflicting attempt for %s", conflict_id)

    return ResolveConflictResponse(
        conflict_id=conflict_id,
        strategy=payload.strategy,
        original_clause=original_clause,
        revised_clause=revised_clause,
        diff=diff,
        reverification=reverification,
        cleared=cleared,
    )


@router.post(
    "/api/conflicts/{conflict_id}/resolve/accept",
    response_model=AcceptConflictResolutionResponse,
    tags=["drafts"],
)
async def accept_conflict_resolution(
    conflict_id: uuid.UUID,
    payload: AcceptConflictResolutionRequest,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> AcceptConflictResolutionResponse:
    """
    Commits a previously-generated revision: patches only the flagged
    clause in the stored draft (the rest of the document and its
    already-passed conflict checks are untouched) and durably marks the
    conflict resolved (resolution_status='resolved' + resolved_clause_text
    persisted, not just returned in this response) so the state survives a
    page reload or a fresh analysis run.
    """
    logger.info("Accept-resolution requested for conflict %s", conflict_id)
    try:
        conflict = await conflicts_repo.get_by_id(session, conflict_id)
    except SQLAlchemyError:
        logger.exception("Database error loading conflict %s", conflict_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    if conflict is None:
        raise HTTPException(status_code=404, detail="Conflict not found")

    draft_row = await _load_row(conflict.generated_draft_id, session)
    _ensure_can_access_draft(draft_row, officer)

    # Idempotent: a second accept call for an already-resolved conflict
    # (e.g. a retry after the caller didn't see the first response) is a
    # no-op success, not an error — the draft content was already patched
    # by the first call and must not be patched again.
    if conflict.resolution_status == "resolved":
        logger.info("Conflict %s already resolved; accept is a no-op", conflict_id)
        return AcceptConflictResolutionResponse(
            conflict=ConflictOut.model_validate(conflict),
            draft_id=draft_row.generated_draft_id,
            draft_version=draft_row.version,
        )

    original_clause = conflict.draft_excerpt or ""
    if not original_clause or original_clause not in draft_row.content:
        try:
            await conflicts_repo.record_resolve_attempt(session, conflict_id, status="attempted_error")
        except SQLAlchemyError:
            logger.exception("Database error recording accept failure for %s", conflict_id)
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="The flagged clause no longer matches the current draft content; it may have "
                   "already been edited. Re-run conflict detection before resolving.",
        )

    new_content = draft_row.content.replace(original_clause, payload.revised_clause, 1)
    new_content_plain = (
        draft_row.content_plain.replace(original_clause, payload.revised_clause, 1)
        if draft_row.content_plain else new_content
    )

    try:
        updated_draft = await drafts_repo.patch_draft_content(
            session,
            conflict.generated_draft_id,
            new_content=new_content,
            new_content_plain=new_content_plain,
            edited_by=officer.officer_id,
            change_note=f"Resolved conflict {conflict.conflict_ref}",
        )
        logger.info("Conflict %s: draft %s content patched (version -> %d)",
                    conflict_id, conflict.generated_draft_id, updated_draft.version)
        updated_conflict = await conflicts_repo.resolve_conflict(
            session, conflict_id,
            reason=f"Resolved via clause revision ({conflict.conflict_ref})",
            resolved_clause_text=payload.revised_clause,
        )
    except drafts_repo.DraftImmutableError as exc:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail=str(exc))
    except SQLAlchemyError:
        logger.exception("Database error accepting conflict resolution %s", conflict_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")

    logger.info("Conflict %s: resolution persisted (resolution_status=resolved)", conflict_id)

    return AcceptConflictResolutionResponse(
        conflict=ConflictOut.model_validate(updated_conflict),
        draft_id=updated_draft.generated_draft_id,
        draft_version=updated_draft.version,
    )


@router.patch("/api/drafts/{draft_id}", response_model=Draft, tags=["drafts"])
async def patch_draft(
    draft_id: uuid.UUID,
    payload: DraftUpdate,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> Draft:
    """
    Save the current editor content (Task 5c). Snapshots the previous
    content into draft_versions before overwriting — see
    db/repositories/drafts.patch_draft_content.
    """
    row = await _load_row(draft_id, session)
    _ensure_can_access_draft(row, officer)

    try:
        updated = await drafts_repo.patch_draft_content(
            session,
            draft_id,
            new_content=payload.body_text,
            new_content_plain=payload.content_plain or payload.body_text,
            edited_by=officer.officer_id,
            change_note=payload.change_note,
        )
    except drafts_repo.DraftImmutableError as exc:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail=str(exc))
    except SQLAlchemyError:
        logger.exception("Database error saving draft %s", draft_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")

    return _to_draft_schema(updated)


@router.delete("/api/drafts/{draft_id}",
               status_code=status.HTTP_204_NO_CONTENT, tags=["drafts"])
async def delete_draft(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> None:
    """Soft delete: sets status = 'archived'. Nothing is ever hard-deleted from here."""
    row = await _load_row(draft_id, session)
    _ensure_can_access_draft(row, officer)
    try:
        await drafts_repo.archive_draft(session, draft_id)
    except SQLAlchemyError:
        logger.exception("Database error archiving draft %s", draft_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")


# =====================================================================
# Analysis — one route per objective, plus a combined one
# =====================================================================

@router.post("/api/analysis/{draft_id}/template",
             response_model=List[TemplateIssue], tags=["analysis"])
async def run_template_check(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> List[TemplateIssue]:
    """Objective 4: Manual of Office Procedure enforcement. Instant, no AI."""
    draft = await _load(draft_id, session, officer)
    if draft.language == Language.MARATHI:
        return template_rules_marathi.check_template_marathi(draft.body_text)
    return template_rules.check_template(draft.body_text)


@router.post("/api/analysis/{draft_id}/references",
             response_model=List[ReferenceHit], tags=["analysis"])
async def run_reference_tracking(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> List[ReferenceHit]:
    """Objective 3: find and resolve every GR this draft cites."""
    draft = await _load(draft_id, session, officer)
    hits = references.extract_references(draft.body_text)
    return references.resolve_against_corpus(hits)


@router.post("/api/analysis/{draft_id}/conflicts",
             response_model=List[ConflictHit], tags=["analysis"])
async def run_conflict_detection(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> List[ConflictHit]:
    """
    Objective 1: cross-departmental conflict detection.

    The slowest route by far, since it makes model calls. The frontend
    should call it separately and show a spinner, rather than blocking
    the whole report on it.

    Delegates to the conflict_detection package (rule engine + LLM verifier) and adapts its
    ConflictReportItem results into the ConflictHit shape the frontend already expects.
    """
    draft = await _load(draft_id, session, officer)
    draft_lang = "mr" if draft.language == Language.MARATHI else "en"

    retrieval_trace: List[ClauseRetrievalTrace] = []
    report_items = detect_cross_department_conflicts(
        draft.body_text, draft_language=draft_lang, trace=retrieval_trace
    )
    report_items = [item for item in report_items
                     if item.confidence >= settings.CONFLICT_CONFIDENCE_FLOOR]

    # Persist the retrieval trace so a "0 conflicts" result stays traceable
    # after the fact (candidates checked-and-cleared vs. never retrieved vs.
    # filtered before the LLM) without needing to re-run detection. Overwrite,
    # not append -- this always reflects the most recent analysis run.
    try:
        draft_row = await _load_row(draft_id, session)
        draft_row.retrieval_trace = [c.model_dump() for c in retrieval_trace]
        await session.flush()
    except SQLAlchemyError:
        logger.exception("Database error persisting retrieval trace for draft %s", draft_id)

    # Persist so each conflict gets a real conflict_id — needed for the
    # Resolve Conflict follow-up action, which operates on a specific ID.
    # Dedupe against this draft's own still-open (unresolved, undismissed)
    # conflicts first: re-running analysis on an unchanged clause would
    # otherwise insert a fresh duplicate row every time.
    conflict_ids: List[Optional[uuid.UUID]] = [None] * len(report_items)
    row_by_id: dict = {}
    existing_rows: list = []
    try:
        existing_rows = await conflicts_repo.get_conflicts_for_draft(session, draft_id, include_dismissed=True)
        row_by_id = {row.conflict_id: row for row in existing_rows}
        existing_by_key = {
            (row.draft_excerpt, row.conflicting_gr_id): row.conflict_id
            for row in existing_rows
            if not row.is_resolved and not row.is_dismissed
        }

        to_insert_items = []
        to_insert_indices = []
        for idx, item in enumerate(report_items):
            key = (item.draft_clause, item.existing_gr_id)
            existing_id = existing_by_key.get(key)
            if existing_id is not None:
                conflict_ids[idx] = existing_id
            else:
                to_insert_items.append(item)
                to_insert_indices.append(idx)

        if to_insert_items:
            persisted_rows = await drafts_repo.persist_conflicts_for_draft(
                session,
                draft_id,
                [
                    {
                        "source_of_conflict": item.existing_department or item.existing_gr_title,
                        "conflicting_text": item.matched_clause,
                        "draft_excerpt": item.draft_clause,
                        "conflicting_gr_id": item.existing_gr_id,
                        "source_gr_title": item.existing_gr_title,
                        "severity": _to_db_severity(item.severity),
                        "justification": item.reason,
                        "detected_by": "llm_verifier",
                    }
                    for item in to_insert_items
                ],
            )
            for idx, row in zip(to_insert_indices, persisted_rows):
                conflict_ids[idx] = row.conflict_id
                row_by_id[row.conflict_id] = row
    except SQLAlchemyError:
        logger.exception("Database error persisting conflicts for draft %s", draft_id)

    live_hits = [
        ConflictHit(
            conflict_id=conflict_id,
            draft_clause=item.draft_clause,
            existing_gr_id=item.existing_gr_id,
            existing_gr_title=item.existing_gr_title,
            existing_department=item.existing_department,
            existing_clause=item.matched_clause,
            relation=Relation.OVERLAP if item.relation == "overlap" else Relation.CONFLICT,
            confidence=item.confidence,
            justification=item.reason,
            source_url=item.source_url,
            conflict_type=item.category,
            severity=item.severity,
            resolution_status=(row_by_id[conflict_id].resolution_status if conflict_id in row_by_id else "not_attempted"),
            resolved_clause_text=(row_by_id[conflict_id].resolved_clause_text if conflict_id in row_by_id else None),
            source_ocr_low_confidence=(row_by_id[conflict_id].source_ocr_low_confidence if conflict_id in row_by_id else False),
        )
        for item, conflict_id in zip(report_items, conflict_ids)
    ]

    # Resolved conflicts stay visible even though a fresh detection pass
    # naturally won't re-surface a clause that no longer conflicts — read
    # from the persisted record rather than recomputing "resolved" from
    # scratch each report load. Anything already covered by live_hits
    # above is skipped so a resolved conflict never appears twice.
    covered_ids = {cid for cid in conflict_ids if cid is not None}
    resolved_hits = [
        ConflictHit(
            conflict_id=row.conflict_id,
            draft_clause=row.draft_excerpt or "",
            existing_gr_id=row.conflicting_gr_id or "",
            existing_gr_title=row.source_gr_title or "",
            existing_department=row.source_of_conflict or "",
            existing_clause=row.conflicting_text or "",
            relation=Relation.OVERLAP,
            confidence=1.0,
            justification=row.justification,
            source_url=None,
            conflict_type="Policy Conflict",
            severity=row.severity.value if hasattr(row.severity, "value") else row.severity,
            resolution_status=row.resolution_status,
            resolved_clause_text=row.resolved_clause_text,
            source_ocr_low_confidence=row.source_ocr_low_confidence,
        )
        for row in existing_rows
        if row.is_resolved and row.conflict_id not in covered_ids
    ]

    # Detection is LLM-based and not perfectly deterministic run-to-run — a
    # clause the OCR upload pipeline (or an earlier analysis pass) already
    # flagged as an open conflict shouldn't vanish just because this
    # particular fresh pass didn't happen to re-surface it. Anything still
    # open (not resolved, not dismissed) and not already covered by
    # live_hits above is carried forward rather than silently dropped.
    stale_open_hits = [
        ConflictHit(
            conflict_id=row.conflict_id,
            draft_clause=row.draft_excerpt or "",
            existing_gr_id=row.conflicting_gr_id or "",
            existing_gr_title=row.source_gr_title or "",
            existing_department=row.source_of_conflict or "",
            existing_clause=row.conflicting_text or "",
            relation=Relation.OVERLAP,
            confidence=1.0,
            justification=row.justification,
            source_url=None,
            conflict_type="Policy Conflict",
            severity=row.severity.value if hasattr(row.severity, "value") else row.severity,
            resolution_status=row.resolution_status,
            resolved_clause_text=row.resolved_clause_text,
            source_ocr_low_confidence=row.source_ocr_low_confidence,
        )
        for row in existing_rows
        if not row.is_resolved and not row.is_dismissed and row.conflict_id not in covered_ids
    ]

    return live_hits + resolved_hits + stale_open_hits


@router.post("/api/analysis/{draft_id}/terminology",
             response_model=List[TermMapping], tags=["analysis"])
async def run_terminology(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> List[TermMapping]:
    """Objective 2: bilingual legal terminology consistency."""
    draft = await _load(draft_id, session, officer)
    return llm.map_terminology(draft.body_text, draft.language)


@router.post("/api/analysis/{draft_id}",
             response_model=AnalysisReport, tags=["analysis"])
async def run_full_analysis(
    draft_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> AnalysisReport:
    """
    Run all four objectives and return one report.
    THIS IS THE ENDPOINT THE MAIN SCREEN CALLS.

    Also persists the detected conflicts and extracted references
    against the draft (Task 1h) the first time it's run — re-running
    analysis on an already-analysed draft does not duplicate rows.
    """
    draft = await _load(draft_id, session, officer)

    if draft.language == Language.MARATHI:
        template_issues = template_rules_marathi.check_template_marathi(draft.body_text)
    else:
        template_issues = template_rules.check_template(draft.body_text)

    reference_hits = references.resolve_against_corpus(
        references.extract_references(draft.body_text)
    )
    conflicts = await run_conflict_detection(draft_id, officer, session)
    terms = llm.map_terminology(draft.body_text, draft.language)

    try:
        if not await drafts_repo.has_analysis_results(session, draft_id):
            await drafts_repo.attach_analysis_results(
                session,
                draft_id,
                conflicts=[_conflict_hit_to_dict(c) for c in conflicts],
                references=[_reference_hit_to_dict(r) for r in reference_hits],
            )
    except SQLAlchemyError:
        # Analysis itself succeeded; don't fail a live demo over a
        # persistence hiccup. The report below is still returned in full.
        logger.exception("Failed to persist analysis results for draft %s", draft_id)

    errors = sum(1 for i in template_issues if i.severity == Severity.ERROR)
    warnings = sum(1 for i in template_issues if i.severity == Severity.WARNING)
    unresolved = sum(1 for r in reference_hits if not r.found_in_corpus)
    top_confidence = max((c.confidence for c in conflicts), default=0.0)

    # A template error or any conflict blocks issue of the GR.
    # A warning or an unresolvable citation only needs a human look.
    if errors or conflicts:
        overall = "blocked"
    elif warnings or unresolved:
        overall = "needs_review"
    else:
        overall = "clean"

    return AnalysisReport(
        draft_id=draft.id,
        generated_at=datetime.now(timezone.utc),
        summary=AnalysisSummary(
            template_error_count=errors,
            template_warning_count=warnings,
            reference_count=len(reference_hits),
            unresolved_reference_count=unresolved,
            conflict_count=len(conflicts),
            highest_conflict_confidence=top_confidence,
            overall_status=overall,
        ),
        template_issues=template_issues,
        references=reference_hits,
        conflicts=conflicts,
        terms=terms,
    )


# =====================================================================
# Corpus search
# =====================================================================

@router.get("/api/corpus/{gr_id}/ocr", response_model=FullOCRResponse, tags=["corpus"])
def get_corpus_ocr(gr_id: str, language: str = Query(None, description="Filter by language 'mr' or 'en'")) -> FullOCRResponse:
    """
    Fetch the full OCR text reconstructed from chunks.
    """
    res = retrieval.get_full_ocr(gr_id, language)
    if not res:
        raise HTTPException(status_code=404, detail=f"GR {gr_id} not found in corpus")
    return FullOCRResponse(**res)


@router.get("/api/official-source/{gr_number}", response_model=OfficialSourceResponse, tags=["corpus"])
def get_official_source(
    gr_number: str,
    department: str = Query(..., description="Department name for lookup routing"),
    date: str = Query(None, description="Optional date string"),
    subject: str = Query(None, description="Optional subject string")
) -> OfficialSourceResponse:
    # 1. Check Cache
    cached = store.get_cached_official_url(gr_number)
    if cached and cached.get("official_url"):
        return OfficialSourceResponse(status="found", url=cached["official_url"])

    # 2. Get Adapter and look up
    adapter = get_adapter(department)
    url = adapter.find_pdf(gr_number, date, subject)

    # 3. Store result in Cache if found
    if url:
        store.set_cached_official_url(gr_number, department, url)
        return OfficialSourceResponse(status="found", url=url)

    return OfficialSourceResponse(status="not_found", url=None)


@router.get("/api/official-gr/{gr_number}", response_model=OfficialSourceResponse, tags=["corpus"])
def get_official_gr(
    gr_number: str,
    department: str = Query(..., description="Department name for lookup routing"),
    date: str = Query(None, description="Optional date string"),
    subject: str = Query(None, description="Optional subject string")
) -> OfficialSourceResponse:
    """
    Locate the exact Government Resolution hosted on the official portal.
    Checks the local cache first, otherwise triggers resolver.
    """
    # 1. Check Cache
    cached = store.get_cached_official_url(gr_number)
    if cached and cached.get("official_url"):
        return OfficialSourceResponse(status="found", url=cached["official_url"])

    # 2. Get Adapter and look up
    from lookup import get_adapter
    adapter = get_adapter(department)
    url = adapter.find_pdf(gr_number, date, subject)

    # 3. Store result in Cache if found
    if url:
        store.set_cached_official_url(gr_number, department, url)
        return OfficialSourceResponse(status="found", url=url)

    return OfficialSourceResponse(status="not_found", url=None)


@router.get("/api/corpus/search",
            response_model=CorpusSearchResponse, tags=["corpus"])
def search_corpus(
    q: str = Query(..., min_length=3, description="Natural language query"),
    top_k: int = Query(default=None, ge=1, le=50),
    officer: Optional[Officer] = Depends(optional_auth),
) -> CorpusSearchResponse:
    """
    Search past GRs by meaning rather than keyword. PUBLIC — no login
    required (Task 2). optional_auth resolves the officer when a token
    is present so the search can be attributed in the logs; results are
    identical either way.
    """
    started = time.perf_counter()
    hits = retrieval.search(q, top_k=top_k or settings.TOP_K)
    took_ms = int((time.perf_counter() - started) * 1000)
    logger.info(
        "Corpus search by %s: %r (%d hits, %d ms)",
        officer.login_id if officer else "anonymous",
        q, len(hits), took_ms,
    )
    return CorpusSearchResponse(query=q, hits=hits, took_ms=took_ms)


@router.get("/api/knowledge/search", tags=["knowledge"])
def search_knowledge_terminology(
    q: str = Query(..., min_length=1, description="Search text in English or Marathi"),
    limit: int = Query(default=10, ge=1, le=50),
):
    """
    Instant terminology lookup in English or Marathi against the Knowledge Base.
    """
    from knowledge import get_knowledge_service
    ks = get_knowledge_service()
    results = ks.search(q, limit=limit)
    return {"query": q, "results": results}


# =====================================================================
# Health
# =====================================================================

@router.get("/api/health/db", response_model=HealthDbResponse, tags=["health"])
async def health_db() -> HealthDbResponse:
    """
    Connectivity + round-trip latency for the local Postgres database.

    Deliberately does NOT use Depends(get_session): that dependency
    commits on cleanup regardless of what the route body did, and if
    Postgres is unreachable, THAT commit throws too — outside this
    function's own try/except, so it reaches the global handler and
    turns a health check into a raw 500 instead of a clean
    connected=False. A short-lived raw connection sidesteps that
    session-commit-on-cleanup entirely.
    """
    started = time.perf_counter()
    try:
        async with db_engine.connect() as conn:
            await conn.execute(text("SELECT 1"))
    except Exception:
        logger.exception("Database health check failed")
        return HealthDbResponse(status="error", connected=False, latency_ms=None)
    latency_ms = (time.perf_counter() - started) * 1000
    return HealthDbResponse(status="ok", connected=True, latency_ms=round(latency_ms, 2))


@router.post("/api/copilot/draft", response_model=DraftGenerateResponse, tags=["copilot"])
async def copilot_draft(
    payload: DraftGenerateRequest,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> DraftGenerateResponse:
    """Generates a draft GR. Requires login (Task 2: "generating a draft")."""
    # 1. Retrieve similar GRs for styling/reference — trim snippets to save tokens
    hits = retrieval.search(payload.prompt, top_k=3)

    # 2. Query KnowledgeService for department & standard phrase guidelines
    from knowledge import get_knowledge_service
    ks = get_knowledge_service()

    def _dept_display(slug: str) -> str:
        entry = ks.find_department(slug)
        if entry:
            return f"{entry.get('english')} ({entry.get('marathi')})"
        return slug.replace("_", " ")

    examples_str = ""
    for idx, hit in enumerate(hits):
        examples_str += f"--- EXAMPLE GR {idx+1} ---\n" f"Department: {_dept_display(hit.department)}\n" f"Title: {hit.title}\n"
        if hit.gr_number:
            examples_str += f"GR Number: {hit.gr_number}\n"
        if hit.issued_on:
            examples_str += f"Date: {hit.issued_on}\n"
        examples_str += f"Excerpt: {hit.snippet[:400]}\n\n"

    dept = payload.department if payload.department else (hits[0].department if hits else "General_Administration_Department")
    dept_display = _dept_display(dept)

    # Select standard phrases from Knowledge Base for mandatory terminology adherence
    std_phrases = ks.get_all_phrases()[:5]
    phrase_guidelines = "\n".join(
        f"  - {p.get('english')} -> {p.get('marathi')}" for p in std_phrases if p.get('english')
    )

    # 3. LLM drafting call
    system_prompt = prompts.COPILOT_DRAFT
    gen_date = prompts.format_generation_date()
    date_str = gen_date["marathi"] if payload.language.lower() == "marathi" else gen_date["english"]
    user_msg = (
        f"Input:\n"
        f"- User Prompt: {payload.prompt}\n"
        f"- Issuing Department: {dept_display}\n"
        f"- Language: {payload.language}\n"
        f"- Generation Date: {date_str}\n"
        f"- Official Standard Phrases:\n{phrase_guidelines}\n\n"
        f"- Retrieved Context:\n{examples_str}"
    )
    body_text = template_rules.strip_llm_formatting_artifacts(llm.call_model(system_prompt, user_msg))

    leaks = template_rules.find_placeholder_leaks(body_text)
    if leaks:
        logger.warning("Draft contained unrendered placeholders %s; retrying once", leaks)
        retry_msg = (
            f"{user_msg}\n\n"
            f"CORRECTION REQUIRED: your previous attempt left these placeholders "
            f"unfilled: {', '.join(leaks)}. Regenerate the full document with every "
            f"placeholder substituted by a real value."
        )
        body_text = template_rules.strip_llm_formatting_artifacts(llm.call_model(system_prompt, retry_msg))
        leaks = template_rules.find_placeholder_leaks(body_text)
        if leaks:
            logger.error("Draft still contains unrendered placeholders after retry: %s", leaks)
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail="Draft generation failed validation: unresolved placeholders in output.",
            )

    title = f"Draft GR: {payload.prompt[:50]}"

    # 4. Persist the draft, attributed to the logged-in officer.
    lang_enum = Language.MARATHI if payload.language.lower() == "marathi" else Language.ENGLISH
    try:
        draft = await drafts_repo.create_draft_with_analysis(
            session,
            title=title,
            language=lang_enum.value,
            drafted_by=officer.officer_id,
            content=body_text,
            content_plain=body_text,
            department=dept,
            brief=payload.prompt,
        )
    except SQLAlchemyError:
        logger.exception("Database error persisting generated draft")
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")

    return DraftGenerateResponse(
        draft_id=str(draft.generated_draft_id),
        title=draft.title,
        department=draft.department,
        body_text=draft.content,
        references=hits,
        gr_number=draft.gr_number,
        language=draft.language.value if hasattr(draft.language, "value") else draft.language,
        status=draft.status.value if hasattr(draft.status, "value") else draft.status,
    )


@router.post("/api/copilot/compare", response_model=ComparisonResponse, tags=["copilot"])
def copilot_compare(payload: ComparisonRequest) -> ComparisonResponse:
    # 1. Look up GR chunks — trim snippets to reduce token usage
    hits1 = retrieval.search(payload.gr_id_1, top_k=3)
    hits2 = retrieval.search(payload.gr_id_2, top_k=3)

    text1 = "\n\n".join(h.snippet[:350] for h in hits1)
    text2 = "\n\n".join(h.snippet[:350] for h in hits2)

    # 2. LLM Comparison
    system_prompt = (
        "You are a policy analyst for the Government of Maharashtra. "
        "Compare the two provided Government Resolutions (GRs) side-by-side. "
        "Create a clean comparative analysis highlighting:\n"
        "1. Eligibility Criteria\n"
        "2. Financial Limits / Allocations\n"
        "3. Departmental Jurisdictions\n"
        "4. Scope of Amendments or applicability\n"
        "Present the output as a clean Markdown table comparing the two, followed by a brief summary of key differences."
    )
    user_msg = f"GOVERNMENT RESOLUTION 1 ({payload.gr_id_1}):\n{text1}\n\nGOVERNMENT RESOLUTION 2 ({payload.gr_id_2}):\n{text2}"
    report = llm.call_model(system_prompt, user_msg)

    return ComparisonResponse(
        gr_id_1=payload.gr_id_1,
        gr_id_2=payload.gr_id_2,
        comparison_report=report
    )


@router.post("/api/copilot/explain-clause", response_model=ClauseExplanationResponse, tags=["copilot"])
def copilot_explain_clause(payload: ClauseExplanationRequest) -> ClauseExplanationResponse:
    system_prompt = (
        "You are a legal advisor. Explain the provided Government Resolution clause in simple terms. "
        "Avoid heavy administrative jargon. Translate or write your explanation in the requested language."
    )
    lang_name = "Marathi" if payload.language == Language.MARATHI else "English"
    user_msg = f"EXPLAIN THIS CLAUSE IN SIMPLE {lang_name}:\n{payload.clause_text}"
    explanation = llm.call_model(system_prompt, user_msg)

    return ClauseExplanationResponse(explanation=explanation)


@router.post("/api/conflicts/detect", response_model=List[ConflictReportItem], tags=["conflicts"])
def detect_conflicts_endpoint(payload: DraftCreate) -> List[ConflictReportItem]:
    """
    Accepts a draft GR and returns a cross-departmental structured conflict report
    using the modular rule-engine and LLM two-stage pipeline.
    """
    draft_lang = "mr" if payload.language == Language.MARATHI else "en"
    return detect_cross_department_conflicts(payload.body_text, draft_language=draft_lang)


MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB

def _extract_text_from_file(filename: str, content: bytes):
    """
    Extract readable text and page count from PDF, DOCX, or TXT content.
    Includes fallback hook for future OCR integrations on scanned PDFs.
    """
    lower_name = filename.lower()
    extracted_text = ""
    page_count = 1

    if lower_name.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            text_pages = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    text_pages.append(txt.strip())
            extracted_text = "\n\n".join(text_pages)
        except Exception as e:
            raise ValueError(f"Failed to read PDF file: {str(e)}")

        # Future OCR hook for scanned PDFs:
        if not extracted_text.strip():
            # Future expansion: run OCR pipeline (tesseract/vision model)
            pass

    elif lower_name.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            extracted_text = "\n\n".join(paragraphs)
            page_count = max(1, len(paragraphs) // 30)
        except Exception as e:
            raise ValueError(f"Failed to read Word document: {str(e)}")

    elif lower_name.endswith(".txt"):
        try:
            extracted_text = content.decode("utf-8")
        except UnicodeDecodeError:
            extracted_text = content.decode("latin-1", errors="ignore")
        page_count = max(1, len(extracted_text.splitlines()) // 40)
    else:
        raise ValueError("Unsupported file format. Only .pdf, .docx, and .txt files are allowed.")

    return extracted_text.strip(), page_count


@router.post("/api/upload-gr/parse-file", tags=["upload"])
async def parse_uploaded_gr_file(file: UploadFile = File(...)):
    """
    Parse uploaded GR file (.pdf, .docx, .txt), validate size and content,
    and return extracted text along with document metadata.
    """
    if not file or not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided."
        )

    filename = file.filename
    lower_name = filename.lower()
    allowed_exts = (".pdf", ".docx", ".txt")
    if not any(lower_name.endswith(ext) for ext in allowed_exts):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file format '{filename}'. Allowed formats: .pdf, .docx, .txt"
        )

    content = await file.read()
    if not content or len(content) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file is empty (0 bytes)."
        )

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds maximum limit of 20 MB."
        )

    ext = lower_name.split(".")[-1]
    try:
        text, page_count = _extract_text_from_file(filename, content)
    except ValueError as val_err:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(val_err)
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal error processing file: {str(exc)}"
        )

    if not text or not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No readable text found in the uploaded document. Please check if the document is password-protected or empty."
        )

    words = text.split()
    word_count = len(words)
    char_count = len(text)

    return {
        "success": True,
        "filename": filename,
        "file_type": ext,
        "page_count": page_count,
        "word_count": word_count,
        "character_count": char_count,
        "text": text
    }


# =====================================================================
# Scanned GR upload — OCR ingestion (images + scanned PDFs)
#
# Distinct from /api/upload-gr/parse-file above: that endpoint reads an
# existing text layer (pypdf) and fails on a scan with none. This path
# assumes there's no text layer at all, OCRs it (Tesseract, eng+mar), and
# routes the result through the *same* detect_cross_department_conflicts()
# typed drafts use — see ocr_ingest/pipeline.py.
# =====================================================================

_OCR_ALLOWED_EXTENSIONS = (".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".tif")


@router.post("/api/gr/upload", response_model=GrUploadResponse, status_code=status.HTTP_202_ACCEPTED, tags=["ocr-upload"])
async def upload_scanned_gr(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> GrUploadResponse:
    """
    Accepts a scanned GR image/PDF, stores it, and schedules OCR + conflict
    detection as a background task — processing can take 30s-3min for a
    multi-page scan, so this returns immediately with a pending record
    rather than blocking the request. Poll GET /api/gr/upload/{upload_id}
    for status.
    """
    if not file or not file.filename:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No file provided.")

    lower_name = file.filename.lower()
    if not any(lower_name.endswith(ext) for ext in _OCR_ALLOWED_EXTENSIONS):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file format '{file.filename}'. Allowed: .pdf, .png, .jpg, .jpeg, .tiff",
        )
    file_type = lower_name.rsplit(".", 1)[-1]
    if file_type == "tif":
        file_type = "tiff"

    content = await file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="The uploaded file is empty (0 bytes).")

    max_bytes = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds the {settings.MAX_UPLOAD_SIZE_MB} MB limit.",
        )

    file_hash = hashlib.sha256(content).hexdigest()

    # Dedup: a byte-identical file already has a record (any status) —
    # return it as-is rather than re-OCRing. If it's still processing, the
    # caller just starts polling this same upload_id.
    try:
        existing = await gr_uploads_repo.get_by_hash(session, file_hash)
    except SQLAlchemyError:
        logger.exception("Database error checking upload dedup for hash %s", file_hash)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    if existing is not None:
        return GrUploadResponse.model_validate(existing)

    try:
        upload = await gr_uploads_repo.create_pending_upload(
            session,
            file_hash=file_hash,
            original_filename=file.filename,
            file_type=file_type,
            file_size_bytes=len(content),
            uploaded_by=officer.officer_id,
        )
    except SQLAlchemyError:
        logger.exception("Database error creating gr_uploads row")
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")

    # Raw bytes go to disk, not Postgres — see ocr_ingest/storage.py.
    save_upload_bytes(file_hash, file_type, content)

    # Explicit commit here, ahead of get_session's usual end-of-request
    # commit: BackgroundTasks can start running before that automatic
    # commit fires (dependency cleanup and background-task execution
    # aren't strictly ordered the way you'd assume), and the task opens
    # its OWN session — it won't see this row at all until it's committed.
    await session.commit()

    background_tasks.add_task(run_ocr_pipeline, upload.upload_id)

    return GrUploadResponse.model_validate(upload)


@router.get("/api/gr/upload/{upload_id}", response_model=GrUploadResponse, tags=["ocr-upload"])
async def get_scanned_gr_upload(
    upload_id: uuid.UUID,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
) -> GrUploadResponse:
    """Poll target for the background OCR job. Once status is
    complete/needs_review, generated_draft_id is set and the frontend
    switches to the ordinary draft endpoints (GET /api/drafts/{id},
    GET /api/drafts/{id}/conflicts, etc.) — unchanged by this feature."""
    try:
        upload = await gr_uploads_repo.get_by_id(session, upload_id)
    except SQLAlchemyError:
        logger.exception("Database error loading gr_upload %s", upload_id)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")
    if upload is None:
        raise HTTPException(status_code=404, detail="Upload not found")

    # Same ownership gate as drafts: uploader, or admin/reviewer.
    if officer.role not in (OfficerRole.ADMIN, OfficerRole.REVIEWER) and upload.uploaded_by != officer.officer_id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="You do not have access to this upload.")

    return GrUploadResponse.model_validate(upload)
