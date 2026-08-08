"""
export_docx.py — Task 5a: renders a stored draft as a .docx formatted
like a real Maharashtra GR.

SQL RULE: the route below never builds a query — it loads the draft via
db.repositories.drafts, the same audited path every other draft read
uses. See backend/README.md, "SQL injection prevention".

DEVANAGARI FONT — the most common failure in this task: Word renders a
run's text using its **complex-script (w:cs) font**, not its ASCII
font, whenever the run contains complex-script code points (Devanagari
included). python-docx's high-level `run.font.name` only ever writes
`w:ascii`/`w:hAnsi`. Leaving `w:cs`/`w:eastAsia` unset falls back to
whatever the style/theme default is (frequently a Latin-only font),
which is exactly how Marathi ends up as empty boxes. `_set_run_font`
below sets all four rFonts slots explicitly and identically, on every
run, regardless of script — see "How this was verified" at the bottom
of this file for how that was checked without a Word install on hand.
"""

import html
import io
import logging
import re
from datetime import date as date_type
from html.parser import HTMLParser
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.ext.asyncio import AsyncSession

from db.base import get_session
from db.models import Officer
from db.repositories import drafts as drafts_repo
from deps import get_current_officer
from schemas import ExportDocxRequest

logger = logging.getLogger("nirn.export_docx")

router = APIRouter()

DEVANAGARI_FONT = "Nirmala UI"
DEVANAGARI_FALLBACK_FONT = "Mangal"


# =====================================================================
# Font handling
# =====================================================================

def _set_run_font(run, *, size_pt: float = 12, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
    """
    Sets ascii, hAnsi, eastAsia AND cs (complex-script) to the same
    Devanagari-capable font on this run, explicitly, via the raw OOXML
    element — not just run.font.name, which only touches ascii/hAnsi
    and leaves w:cs (what Word actually uses to render Devanagari)
    untouched. Mangal is recorded as the documented fallback in the
    document's default styles (see _configure_fallback_style) for
    machines without Nirmala UI installed.
    """
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(slot), DEVANAGARI_FONT)


def _configure_fallback_style(document: Document) -> None:
    """
    Sets Mangal as the complex-script font on the document's Normal
    style, so any text Word lays out from style defaults (rather than
    an explicit per-run override) still resolves to a Devanagari-capable
    font on machines where Nirmala UI isn't installed.
    """
    normal = document.styles["Normal"]
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), DEVANAGARI_FONT)
    rFonts.set(qn("w:hAnsi"), DEVANAGARI_FONT)
    rFonts.set(qn("w:eastAsia"), DEVANAGARI_FONT)
    rFonts.set(qn("w:cs"), DEVANAGARI_FALLBACK_FONT)


# =====================================================================
# Tiptap HTML -> docx mapping
#
# Never dumps raw HTML tags into the document — walks the tag tree and
# maps each element to the matching Word construct.
# =====================================================================

class _DocxHtmlBuilder(HTMLParser):
    _HEADING_SIZES = {"h1": 20, "h2": 16, "h3": 14}

    def __init__(self, document: Document):
        super().__init__()
        self.document = document
        self._bold = False
        self._italic = False
        self._underline = False
        self._list_stack: List[str] = []  # "ul" | "ol"
        self._current_paragraph = None
        self._current_style = "body"  # "body" | "h1" | "h2" | "h3" | "list"

    def _start_paragraph(self, style: str = "body") -> None:
        self._current_style = style
        p = self.document.add_paragraph()
        if style == "body":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif style == "list_bullet":
            p.style = "List Bullet"
        elif style == "list_number":
            p.style = "List Number"
        elif style in ("h1", "h2", "h3"):
            # Word's actual Heading N paragraph style — not just a bold,
            # larger run — so it participates in the outline/navigation
            # pane and any Word-generated table of contents.
            p.style = f"Heading {style[1]}"
        self._current_paragraph = p

    def _ensure_paragraph(self):
        if self._current_paragraph is None:
            self._start_paragraph("body")
        return self._current_paragraph

    def handle_starttag(self, tag, attrs):
        if tag in ("h1", "h2", "h3"):
            self._start_paragraph(tag)
        elif tag == "p":
            self._start_paragraph("body")
        elif tag == "ul":
            self._list_stack.append("ul")
        elif tag == "ol":
            self._list_stack.append("ol")
        elif tag == "li":
            list_type = self._list_stack[-1] if self._list_stack else "ul"
            self._start_paragraph("list_bullet" if list_type == "ul" else "list_number")
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "u":
            self._underline = True
        elif tag == "br":
            self._ensure_paragraph().add_run().add_break()

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "p", "li"):
            self._current_paragraph = None
        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "u":
            self._underline = False

    def handle_data(self, data):
        if not data or not data.strip():
            if data != " ":
                return
        p = self._ensure_paragraph()
        run = p.add_run(data)
        size = self._HEADING_SIZES.get(self._current_style, 12)
        heading_bold = self._current_style in ("h1", "h2", "h3")
        _set_run_font(
            run,
            size_pt=size,
            bold=self._bold or heading_bold,
            italic=self._italic,
            underline=self._underline,
        )


_HTML_START_RE = re.compile(r"^<(p|div|h[1-6]|table)[\s>]", re.IGNORECASE)


def _looks_like_html(text: str) -> bool:
    return bool(_HTML_START_RE.match((text or "").strip()))


def _plain_text_to_paragraph_html(text: str) -> str:
    """
    A draft's `content` is raw LLM plain text (with literal '\\n' line
    breaks) until an officer explicitly saves an edit through the
    Tiptap editor, which submits real HTML instead. Feeding that plain
    text straight into _DocxHtmlBuilder used to produce ONE giant
    justified paragraph, with python-docx silently turning every
    embedded '\\n' into a <w:br/> soft line break rather than a new
    paragraph. Word then justifies each of those short forced lines
    individually — stretching a five-word line to fill the entire
    text width — which is what "text scattered all over the page"
    actually was. Splitting on blank lines into real paragraphs (no
    embedded '\\n' ever reaches add_run) avoids that entirely.
    """
    lines = (text or "").replace("\r\n", "\n").split("\n")
    paragraphs: List[str] = []
    buffer: List[str] = []

    def flush() -> None:
        if buffer:
            merged = " ".join(buffer).strip()
            if merged:
                paragraphs.append(merged)
            buffer.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush()
            continue
        buffer.append(line)
    flush()

    if not paragraphs:
        return "<p></p>"
    return "".join(f"<p>{html.escape(p)}</p>" for p in paragraphs)


_GOVERNOR_PHRASE = "महाराष्ट्राचे राज्यपाल यांच्या आदेशानुसार व नावाने"


def _strip_llm_closing(content_html: str) -> str:
    """
    The generation prompt reliably gets the LLM to write its own closing
    (governor-order phrase + signature + प्रत list) as the tail of the
    draft body — this phrase is the fixed, never-varies idiom real GRs
    always close with, so its first occurrence is always the start of
    that closing block, never mid-document body text. _add_signature_block
    and _add_distribution_list below add a second, correct one (real
    officer name instead of the LLM's placeholder, full distribution
    list instead of a short guessed one) — without this, both end up in
    the exported document, back to back.
    """
    idx = content_html.find(_GOVERNOR_PHRASE)
    if idx == -1:
        return content_html
    tag_start = content_html.rfind("<", 0, idx)
    return content_html[:tag_start] if tag_start != -1 else content_html[:idx]


def _render_html_body(document: Document, html_content: str) -> None:
    content = html_content or ""
    if not _looks_like_html(content):
        content = _plain_text_to_paragraph_html(content)
    content = _strip_llm_closing(content)
    builder = _DocxHtmlBuilder(document)
    builder.feed(content)
    builder.close()


# =====================================================================
# Header / subject / references / signature blocks
# =====================================================================

def _add_centered_run(paragraph, text_, *, size_pt=12, bold=False):
    run = paragraph.add_run(text_)
    _set_run_font(run, size_pt=size_pt, bold=bold)
    return run


def _add_header_block(document: Document, *, department: str, gr_number: Optional[str], issue_date: str, is_marathi: bool) -> None:
    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered_run(p1, "महाराष्ट्र शासन", size_pt=16, bold=True)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered_run(p3, department.replace("_", " "), size_pt=12, bold=True)

    gr_label = "शासन निर्णय क्रमांक" if is_marathi else "Government Resolution No."
    date_label = "दिनांक" if is_marathi else "Date"

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered_run(p4, f"{gr_label}: {gr_number or 'PROVISIONAL — NOT YET ISSUED'}", size_pt=11)

    p5 = document.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered_run(p5, "मंत्रालय, मुंबई- 400 032", size_pt=11)

    p6 = document.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_centered_run(p6, f"{date_label}: {issue_date}", size_pt=11)

    # Horizontal rule beneath the header, via a bottom paragraph border.
    hr_p = document.add_paragraph()
    p_pr = hr_p._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _add_subject(document: Document, title: str, is_marathi: bool) -> None:
    p = document.add_paragraph()
    label = "विषय" if is_marathi else "Subject"
    run_label = p.add_run(f"{label}: ")
    _set_run_font(run_label, size_pt=12, bold=True)
    run_title = p.add_run(title)
    _set_run_font(run_title, size_pt=12, bold=True)


def _add_references_section(document: Document, references, is_marathi: bool) -> None:
    if not references:
        return
    heading = document.add_paragraph()
    run = heading.add_run("वाचा:-")
    _set_run_font(run, size_pt=12, bold=True)

    for ref in references:
        p = document.add_paragraph(style="List Number")
        text_ = ref.reference_text if hasattr(ref, "reference_text") else str(ref)
        run = p.add_run(text_)
        _set_run_font(run, size_pt=12)


def _add_signature_block(document: Document, *, name: str, designation: str, is_marathi: bool) -> None:
    document.add_paragraph()  # spacer
    p0 = document.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_centered_run(p0, "महाराष्ट्राचे राज्यपाल यांच्या आदेशानुसार व नावाने.", size_pt=12)

    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_centered_run(p1, f"({name})", size_pt=12, bold=True)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_centered_run(p2, designation or ("शासकीय अधिकारी" if is_marathi else "Government Officer"), size_pt=12)


_DISTRIBUTION_LIST = [
    "मा. राज्यपाल यांचे सचिव.",
    "मा.मुख्यमंत्री यांचे सचिव.",
    "मा. मंत्री (सर्व) यांचे खाजगी सचिव.",
    "मुख्य सचिव, महाराष्ट्र राज्य.",
    "सर्व अपर मुख्य सचिव/प्रधान सचिव, मंत्रालय, मुंबई.",
    "सर्व विभागीय आयुक्त.",
    "सर्व जिल्हाधिकारी.",
    "महालेखापाल एक व दोन, महाराष्ट्र राज्य, मुंबई/नागपूर.",
    "सर्व मंत्रालयीन विभाग.",
    "निवड नस्ती.",
]


def _add_distribution_list(document: Document, is_marathi: bool) -> None:
    """
    A fixed, static प्रत (distribution list) — every real GR ends with one,
    and this system has no per-draft distribution data to compute one from.
    Always rendered in Marathi regardless of the draft's own language,
    matching what real GRs actually do.
    """
    document.add_paragraph()  # spacer
    heading = document.add_paragraph()
    run = heading.add_run("प्रत,")
    _set_run_font(run, size_pt=12, bold=True)

    for item in _DISTRIBUTION_LIST:
        p = document.add_paragraph(style="List Number")
        run = p.add_run(item)
        _set_run_font(run, size_pt=12)


def _add_page_numbers(document: Document) -> None:
    """Page numbers in the footer: `PAGE` field, centered."""
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    _set_run_font(run, size_pt=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# =====================================================================
# Entry point
# =====================================================================

def generate_gr_docx(
    *,
    title: str,
    department: str,
    gr_number: Optional[str],
    issue_date: date_type,
    content_html: str,
    references,
    officer_name: str,
    officer_designation: Optional[str],
    is_marathi: bool,
) -> bytes:
    document = Document()

    section = document.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

    _configure_fallback_style(document)

    _add_header_block(
        document,
        department=department,
        gr_number=gr_number,
        issue_date=issue_date.strftime("%d/%m/%Y"),
        is_marathi=is_marathi,
    )
    _add_subject(document, title, is_marathi)
    _add_references_section(document, references, is_marathi)
    _render_html_body(document, content_html)
    _add_signature_block(document, name=officer_name, designation=officer_designation, is_marathi=is_marathi)
    _add_distribution_list(document, is_marathi)
    _add_page_numbers(document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# =====================================================================
# Route
# =====================================================================

@router.post("/api/export/docx", tags=["export"])
async def export_docx(
    payload: ExportDocxRequest,
    officer: Officer = Depends(get_current_officer),
    session: AsyncSession = Depends(get_session),
):
    """
    Content and metadata are pulled server-side from the draft record —
    never trust client-supplied content for what gets stamped as an
    official-looking document. Wrapped end-to-end in try/except: this
    runs live in a demo and must never 500 with a raw traceback.
    """
    try:
        row = await drafts_repo.get_draft_by_id(session, payload.draft_id)
    except SQLAlchemyError:
        logger.exception("Database error loading draft for export")
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Service unavailable.")

    if row is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Draft not found.")
    if officer.role.value not in ("admin", "reviewer") and row.drafted_by != officer.officer_id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="You do not have access to this draft.")

    is_marathi = (row.language.value if hasattr(row.language, "value") else row.language) == "mr"

    try:
        docx_bytes = generate_gr_docx(
            title=row.title,
            department=row.department,
            gr_number=row.gr_number,
            issue_date=row.created_at.date(),
            content_html=row.content,
            references=row.references,
            officer_name=row.officer.name if row.officer else officer.name,
            officer_designation=(row.officer.designation if row.officer else officer.designation),
            is_marathi=is_marathi,
        )
    except Exception:
        logger.exception("DOCX generation failed for draft %s", payload.draft_id)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Could not generate the document. Please try again.")

    date_str = row.created_at.strftime("%Y%m%d")
    filename = f"GR_{row.gr_number or row.generated_draft_id}_{date_str}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =====================================================================
# How this was verified
#
# No copy of Microsoft Word is available in this environment (macOS,
# headless), so rendering fidelity was checked structurally instead of
# visually:
#   1. Generated a .docx from a draft containing Marathi body text.
#   2. Unzipped it and inspected word/document.xml directly: every
#      <w:r> produced by _DocxHtmlBuilder carries a <w:rFonts> child
#      with w:ascii, w:hAnsi, w:eastAsia AND w:cs all set to
#      "Nirmala UI" (the ONE most common cause of Devanagari-as-boxes
#      is w:cs being left at the style default, since Word renders
#      complex-script code points — Devanagari included — using w:cs,
#      not w:ascii).
#   3. Re-opened the generated file with python-docx and confirmed
#      paragraph.text round-trips the original Devanagari characters
#      byte-for-byte (rules out a mis-encoding turning the glyphs into
#      "?" or replacement characters before they ever reach Word).
#   4. Confirmed styles.xml's Normal style also carries explicit
#      rFonts (ascii/hAnsi/eastAsia = Nirmala UI, cs = Mangal) so any
#      text laid out from a style default rather than an explicit
#      per-run override still resolves to a Devanagari-capable font.
# This does not substitute for opening the file in Word — flagged
# explicitly in the final report as unverified by direct visual
# inspection.
# =====================================================================
